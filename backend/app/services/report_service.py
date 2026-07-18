"""
Report generation service.
Creates downloadable DOCX reports using python-docx and PDF reports using ReportLab.
"""

from __future__ import annotations

from datetime import datetime
import os
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Use ReportLab for clean, zero-dependency PDF rendering on Windows
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from app.core.config import settings

REPORTS_DIR = Path(settings.UPLOAD_DIR) / "reports"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

def _md_to_html(text: str) -> str:
    import re
    # Remove headers that are redundant
    text = re.sub(r'#+\s*(Executive Summary|Key Metrics|Primary Insights|Data Quality & Anomalies)', '', text)
    # Convert remaining headers to bold
    text = re.sub(r'#+\s*(.*)', r'<b>\1</b>', text)
    # Convert bold markdown to bold HTML tags
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    # Convert newlines to breaks
    text = text.replace('\n', '<br/>')
    # Replace bullet point characters if they are there
    text = text.replace('•', '&bull;')
    return text

class ReportService:
    def __init__(self) -> None:
        pass

    def generate_pdf(
        self,
        analysis_id: int,
        dataset_name: str,
        kpis: list[dict[str, Any]],
        insights: list[dict[str, Any]],
        statistics: list[dict[str, Any]],
        summary: str,
        recommendations: list[str],
    ) -> str:
        """Generate PDF report using ReportLab and return the absolute file path."""
        file_name = f"report_{analysis_id}_{int(datetime.now().timestamp())}.pdf"
        output_path = REPORTS_DIR / file_name

        # Page margins
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=45,
            leftMargin=45,
            topMargin=45,
            bottomMargin=45
        )
        
        styles = getSampleStyleSheet()
        
        # Color definitions matching brand Sand & Charcoal system
        brand_charcoal = colors.HexColor('#1D1A18')
        brand_sand = colors.HexColor('#8A8380')
        dark_text = colors.HexColor('#1D1A18')
        light_border = colors.HexColor('#E4E4E7')
        gray_bg = colors.HexColor('#F5F5F5')
        zinc_500 = colors.HexColor('#71717a')

        # Custom Paragraph styles
        title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=22,
            leading=26,
            textColor=brand_charcoal,
            alignment=1, # Center
            spaceAfter=10
        )
        
        meta_style = ParagraphStyle(
            'ReportMeta',
            parent=styles['Italic'],
            fontName='Helvetica-Oblique',
            fontSize=10,
            leading=13,
            textColor=zinc_500,
            alignment=1, # Center
            spaceAfter=25
        )

        h1_style = ParagraphStyle(
            'ReportH1',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=13,
            leading=17,
            textColor=brand_charcoal,
            spaceBefore=18,
            spaceAfter=8,
            keepWithNext=True
        )

        body_style = ParagraphStyle(
            'ReportBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#27272a'),
            spaceAfter=8
        )

        table_header_style = ParagraphStyle(
            'TableHeader',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=9,
            leading=11,
            textColor=colors.white
        )

        table_cell_style = ParagraphStyle(
            'TableCell',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9,
            leading=11,
            textColor=colors.HexColor('#27272a')
        )

        story = []

        # Cover Title
        story.append(Spacer(1, 15))
        story.append(Paragraph("DetectiveAI Business Report", title_style))
        story.append(Paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} &nbsp;|&nbsp; "
            f"Dataset: {dataset_name} &nbsp;|&nbsp; "
            f"Analysis Case: #{analysis_id}",
            meta_style
        ))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", h1_style))
        formatted_summary = _md_to_html(summary) if summary else "No executive summary compiled for this case."
        story.append(Paragraph(formatted_summary, body_style))
        story.append(Spacer(1, 10))

        # KPIs Section
        story.append(Paragraph("Key Performance Indicators", h1_style))
        if kpis:
            kpi_data = [[
                Paragraph("KPI Name", table_header_style),
                Paragraph("Current Value", table_header_style),
                Paragraph("Trend", table_header_style),
                Paragraph("Change %", table_header_style)
            ]]
            for k in kpis:
                kpi_data.append([
                    Paragraph(str(k.get("name", "")), table_cell_style),
                    Paragraph(str(k.get("formatted_value", "")), table_cell_style),
                    Paragraph(str(k.get("trend", "")).upper(), table_cell_style),
                    Paragraph(f"{k.get('change_percentage') or 0.0}%", table_cell_style)
                ])
            
            kpi_table = Table(kpi_data, colWidths=[200, 100, 100, 120])
            kpi_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), brand_charcoal),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, light_border),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, gray_bg]),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ]))
            story.append(kpi_table)
        else:
            story.append(Paragraph("No major KPIs detected in this dataset.", body_style))
        story.append(Spacer(1, 10))

        # Discovered Insights
        story.append(Paragraph("Discovered Insights", h1_style))
        if insights:
            for idx, i in enumerate(insights):
                bullet_text = f"<b>[{i.get('category', '').upper()}]</b> {i.get('description', '')} (Confidence: {i.get('confidence_score', 0)}%)"
                story.append(Paragraph(f"&bull; &nbsp; {bullet_text}", body_style))
        else:
            story.append(Paragraph("No critical anomalies or patterns discovered.", body_style))
        story.append(Spacer(1, 10))

        # Statistical Summary Table
        story.append(Paragraph("Statistical Analysis Summary", h1_style))
        if statistics:
            stat_data = [[
                Paragraph("Test Name", table_header_style),
                Paragraph("Statistic", table_header_style),
                Paragraph("P-Value", table_header_style),
                Paragraph("Significant?", table_header_style)
            ]]
            for s in statistics:
                sig_text = "YES" if s.get("significant") else "NO"
                stat_data.append([
                    Paragraph(str(s.get("test_name", "")), table_cell_style),
                    Paragraph(f"{s.get('statistic', 0.0):.4f}", table_cell_style),
                    Paragraph(f"{s.get('p_value', 1.0):.4f}", table_cell_style),
                    Paragraph(sig_text, table_cell_style)
                ])
            
            stat_table = Table(stat_data, colWidths=[220, 100, 100, 100])
            stat_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), brand_charcoal),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, light_border),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, gray_bg]),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ]))
            story.append(stat_table)
        else:
            story.append(Paragraph("No analytical statistical models run.", body_style))
        story.append(Spacer(1, 10))

        # Recommendations
        story.append(Paragraph("Recommendations & Remediation", h1_style))
        if recommendations:
            for idx, r in enumerate(recommendations):
                story.append(Paragraph(f"{idx+1}. &nbsp; {r}", body_style))
        else:
            story.append(Paragraph("No actionable suggestions computed.", body_style))

        # Build document
        doc.build(story)
        return str(output_path)

    def generate_docx(
        self,
        analysis_id: int,
        dataset_name: str,
        kpis: list[dict[str, Any]],
        insights: list[dict[str, Any]],
        statistics: list[dict[str, Any]],
        summary: str,
        recommendations: list[str],
    ) -> str:
        """Generate DOCX report using python-docx and return the absolute file path."""
        doc = Document()
        
        # Style Title
        title = doc.add_heading("DetectiveAI Business Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle/Meta
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | "
            f"Dataset: {dataset_name} | "
            f"Analysis ID: #{analysis_id}"
        )
        run.font.size = Pt(10)
        run.font.italic = True
        
        # Summary Box
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(summary or "No executive summary compiled for this case.")
        
        # KPIs Section
        doc.add_heading("Key Performance Indicators (KPIs)", level=1)
        if kpis:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Shading Accent 1"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "KPI Name"
            hdr_cells[1].text = "Current Value"
            hdr_cells[2].text = "Trend"
            hdr_cells[3].text = "Change %"

            for k in kpis:
                row_cells = table.add_row().cells
                row_cells[0].text = str(k.get("name", ""))
                row_cells[1].text = str(k.get("formatted_value", ""))
                row_cells[2].text = str(k.get("trend", "")).upper()
                row_cells[3].text = f"{k.get('change_percentage') or 0.0}%"
        else:
            doc.add_paragraph("No major KPIs detected.")

        # Insights Section
        doc.add_heading("Discovered Insights", level=1)
        if insights:
            for idx, i in enumerate(insights):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                run_cat = p.add_run(f"[{i.get('category', '').upper()}] ")
                run_cat.bold = True
                p.add_run(f"{i.get('description', '')} (Confidence: {i.get('confidence_score', 0)}%)")
        else:
            doc.add_paragraph("No major insights discovered.")

        # Statistical Summary Section
        doc.add_heading("Statistical Summary", level=1)
        if statistics:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Shading Accent 1"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Test Name"
            hdr_cells[1].text = "Statistic"
            hdr_cells[2].text = "P-Value"
            hdr_cells[3].text = "Significant?"

            for s in statistics:
                row_cells = table.add_row().cells
                row_cells[0].text = str(s.get("test_name", ""))
                row_cells[1].text = f"{s.get('statistic', 0.0):.4f}"
                row_cells[2].text = f"{s.get('p_value', 1.0):.4f}"
                row_cells[3].text = "YES" if s.get("significant") else "NO"
        else:
            doc.add_paragraph("No statistical tests executed.")

        # Recommendations Section
        doc.add_heading("Recommendations", level=1)
        if recommendations:
            for idx, r in enumerate(recommendations):
                doc.add_paragraph(f"{idx+1}. {r}")
        else:
            doc.add_paragraph("No actionable suggestions computed.")

        file_name = f"report_{analysis_id}_{int(datetime.now().timestamp())}.docx"
        output_path = REPORTS_DIR / file_name
        doc.save(str(output_path))
        return str(output_path)
